"""
resume.py — Resume parser. PDF + DOCX → structured data → memory.
"""

import os, json, re
from pathlib import Path


def extract_text_pdf(path: str) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t: parts.append(t)
    return "\n".join(parts)


def extract_text_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text: parts.append(row_text)
    return "\n".join(parts)


def extract_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":             return extract_text_pdf(path)
    elif ext in (".docx", ".doc"):return extract_text_docx(path)
    else: raise ValueError(f"Unsupported format: {ext}. Use PDF or DOCX.")


def parse_with_ai(raw_text: str, filename: str) -> dict:
    from core.ai_router import ask
    truncated = raw_text[:8000]
    prompt = f"""Parse this resume. Return ONLY a JSON object:
{{
  "name": "Full Name",
  "current_role": "Most recent job title",
  "current_company": "Most recent employer",
  "years_exp": 4.5,
  "target_role": "Role this person is targeting (infer from experience)",
  "skills": ["Python", "SQL"],
  "education": [{{"degree":"...", "institution":"...", "year":"..."}}],
  "past_companies": ["Company A"],
  "past_roles": ["Role A"],
  "projects": [{{"name":"...", "description":"...", "tech":["..."]}}],
  "summary": "2-3 sentence professional summary"
}}

Resume:
{truncated}

ONLY the JSON. No markdown."""

    result = ask(prompt)
    text = result["text"].strip().lstrip("```json").lstrip("```").rstrip("```").strip()
    try:
        data = json.loads(text)
        data["filename"] = filename
        data["raw_text"] = raw_text
        return data
    except Exception:
        return {
            "filename": filename, "name": _name_heuristic(raw_text),
            "current_role": None, "current_company": None,
            "years_exp": _years_heuristic(raw_text), "target_role": None,
            "skills": _skills_heuristic(raw_text),
            "education": [], "past_companies": [], "past_roles": [], "projects": [],
            "summary": None, "raw_text": raw_text,
        }


def parse_resume(file_path: str) -> dict:
    from core.memory import save_resume, log_event
    path = Path(file_path)
    if not path.exists(): raise FileNotFoundError(f"Not found: {file_path}")
    raw = extract_text(str(path))
    if len(raw.strip()) < 50: raise ValueError("Could not extract text from resume.")
    data = parse_with_ai(raw, path.name)
    save_resume(data)
    log_event("resume_parsed", f"Parsed: {path.name}",
              detail=f"{data.get('name')} | {data.get('current_role')}")
    return data


_SKILLS = {
    "python","java","javascript","typescript","go","rust","c++","c#","ruby",
    "swift","kotlin","scala","react","angular","vue","node","django","flask",
    "fastapi","spring","tensorflow","pytorch","pandas","numpy","sql","mysql",
    "postgresql","mongodb","redis","aws","gcp","azure","docker","kubernetes",
    "terraform","git","linux","kafka","spark","machine learning","deep learning",
}

def _skills_heuristic(text):
    low = text.lower()
    return sorted({s.title() for s in _SKILLS if s in low})

def _name_heuristic(text):
    for line in text.split("\n"):
        line = line.strip()
        if line and len(line.split()) <= 4 and not any(c.isdigit() for c in line):
            return line
    return ""

def _years_heuristic(text):
    years = [int(y) for y in re.findall(r'\b(20\d{2}|19\d{2})\b', text)]
    return round(min(max(years) - min(years), 25), 1) if len(years) >= 2 else 0.0
